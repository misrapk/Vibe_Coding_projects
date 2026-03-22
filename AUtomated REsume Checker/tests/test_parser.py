import docx
import os
from src.backend.parser import ResumeParser
import json

def create_test_resume(path):
    doc = docx.Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('john.doe@email.com | +1-123-456-7890 | linkedin.com/in/johndoe')
    
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Software Engineer at Tech Corp (2020 - Present)')
    doc.add_paragraph('Developed web applications using Python and React. Managed AWS infrastructure.')
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph('Bachelor of Science in Computer Science, University of Technology (2016-2020)')
    doc.add_paragraph('GPA: 3.8/4.0')
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, JavaScript, React, Django, AWS, Docker, SQL')
    
    doc.save(path)

if __name__ == "__main__":
    test_path = "test_resume.docx"
    create_test_resume(test_path)
    
    parser = ResumeParser()
    result = parser.parse(test_path)
    
    print(json.dumps(result, indent=2))
    
    # Clean up
    if os.path.exists(test_path):
        os.remove(test_path)
