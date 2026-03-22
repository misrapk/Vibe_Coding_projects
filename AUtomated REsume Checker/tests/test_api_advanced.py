import sys
import os
# Use a test database
TEST_DB = "data/test_features.db"
os.environ["DATABASE_PATH"] = TEST_DB

sys.path.append(os.getcwd())
import pytest
from fastapi.testclient import TestClient
from src.backend.api.main import app
from src.backend.database import DatabaseManager
from src.backend.models import Base, User, Job, Resume, Match, Feedback, UserRole
import shutil
client = TestClient(app)

@pytest.fixture(scope="module", autouse=True)
def setup_db():
    # Setup
    if os.path.exists(TEST_DB):
        try:
            os.remove(TEST_DB)
        except:
            pass
    
    # Force use test DB in models/database by env var
    os.environ["DATABASE_PATH"] = TEST_DB
    
    db = DatabaseManager(TEST_DB)
    # Ensure tables are fresh
    Base.metadata.drop_all(bind=db.engine)
    Base.metadata.create_all(bind=db.engine)
    
    yield db
    # Teardown
    db.close()
    if os.path.exists(TEST_DB):
        try:
            os.remove(TEST_DB)
        except:
            pass

def test_full_recruitment_flow():
    # 1. Register Users
    recruiter_data = {
        "email": "recruiter_adv@example.com",
        "password": "password123",
        "first_name": "Adv",
        "last_name": "Recruiter",
        "role": "recruiter"
    }
    candidate_data = {
        "email": "candidate_adv@example.com",
        "password": "password123",
        "first_name": "Adv",
        "last_name": "Candidate",
        "role": "candidate"
    }
    
    client.post("/api/auth/register", json=recruiter_data)
    client.post("/api/auth/register", json=candidate_data)
    
    # 2. Login
    resp_r = client.post("/api/auth/login", data={"username": recruiter_data["email"], "password": recruiter_data["password"]})
    token_r = resp_r.json()["access_token"]
    headers_r = {"Authorization": f"Bearer {token_r}"}
    
    resp_c = client.post("/api/auth/login", data={"username": candidate_data["email"], "password": candidate_data["password"]})
    token_c = resp_c.json()["access_token"]
    headers_c = {"Authorization": f"Bearer {token_c}"}
    
    # 3. Recruiter Posts Job
    job_payload = {
        "title": "Senior Python Developer",
        "company_name": "AI Tech Corp",
        "description": "We are looking for a Senior Python Developer with experience in FastAPI, AWS, and Machine Learning. 5+ years of experience required.",
        "location": "Remote",
        "salary_min": 120000,
        "salary_max": 180000
    }
    resp_job = client.post("/api/jobs/", json=job_payload, headers=headers_r)
    assert resp_job.status_code == 200
    job_id = resp_job.json()["id"]
    
    # 4. Candidate Uploads Resume
    from docx import Document
    doc = Document()
    doc.add_heading('Adv Candidate', 0)
    doc.add_paragraph('Python Developer with 5 years of experience in FastAPI and Machine Learning.')
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, FastAPI, Machine Learning, SQL, Docker')
    doc.save("dummy_resume.docx")
    
    with open("dummy_resume.docx", "rb") as f:
        client.post("/api/resumes/upload", files={"file": ("resume.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}, headers=headers_c)
    os.remove("dummy_resume.docx")
    
    # 5. Trigger Matching
    resp_match_trigger = client.post(f"/api/matches/trigger/{job_id}", headers=headers_r)
    assert resp_match_trigger.status_code == 200
    assert resp_match_trigger.json()["job_started"] is True
    
    # Sleep to allow background task to run (since it's a test client, 
    # it might run in the same thread but let's be safe)
    import time
    time.sleep(2)
    
    # 6. Check Matches
    resp_candidates = client.get(f"/api/jobs/{job_id}/candidates", headers=headers_r)
    candidates = resp_candidates.json()
    assert len(candidates) > 0
    match_id = 1 # Assuming it's the first one
    
    # 7. Provide Feedback
    feedback_payload = {"action": "shortlist", "notes": "Impressive background in Python."}
    resp_feedback = client.post(f"/api/matches/{match_id}/feedback", json=feedback_payload, headers=headers_r)
    assert resp_feedback.status_code == 200
    
    # 8. Check Analytics
    resp_ana_r = client.get("/api/analytics/dashboard", headers=headers_r)
    assert resp_ana_r.status_code == 200
    assert resp_ana_r.json()["stats"]["total_jobs_posted"] == 1
    assert resp_ana_r.json()["stats"]["shortlisted_candidates"] == 1
    
    resp_ana_c = client.get("/api/analytics/dashboard", headers=headers_c)
    assert resp_ana_c.status_code == 200
    assert resp_ana_c.json()["stats"]["total_resumes"] == 1
    
    # 10. Health & Skills
    assert client.get("/api/health").status_code == 200
    resp_skills = client.get("/api/skills/search?q=python")
    assert any("Python" in s["skill_name"] for s in resp_skills.json())
